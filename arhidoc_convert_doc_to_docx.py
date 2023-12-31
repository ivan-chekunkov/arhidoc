from pathlib import Path

import docx
from doc2docx import convert as doc2docx_convert


def _convert_doc_to_docx(in_path: Path, out_path: Path) -> None:
    """Конвертация из doc в docx при помощи библиотеки doc2docx"""
    try:
        if in_path.suffix.upper() != '.DOC':
            print('Не корректное расширение файла {}'.format(in_path.suffix))
            return
        file_name: str = in_path.name.split('.')[0]
        new_file: Path = Path(out_path).joinpath(file_name + '.docx')
        print('Обрабатываю файл {}!'.format(file_name))
        doc2docx_convert(input_path=in_path, output_path=new_file)
    except Exception as error:
        print(error)


def _iterdir(path: Path) -> list[Path]:
    """Получение списка файлов"""
    return [x for x in path.iterdir() if x.is_file()]


def convert_doc_to_docx() -> None:
    """Конвертация файлов из doc в docx"""
    files = _iterdir(Path('input'))
    index: int = 0
    len_files = len(files)
    for path_file in files:
        index += 1
        print('Обрабатываю {} из {}!'.format(index, len_files))
        _convert_doc_to_docx(path_file, Path('output'))


def read_csv(path_file: Path) -> dict[int, tuple[str, str, str]]:
    """Создание словаря документов с их параметрами на основе выгрузки из базы"""
    with open(path_file, 'r', encoding='cp1251') as file:
        lines = [line.rstrip('\n').split('***') for line in file.readlines()]
    doc_info: dict[int, tuple[str, str, str]] = {}
    for line in lines:
        if len(line) != 4:
            print('Error {}'.format(line))
            continue
        data: str = line[0]
        number: str = line[1].replace('"', '')
        name: str = line[2].replace('"', '')
        index: int = int(line[3])
        if doc_info.get(index, False):
            print('Index double {}'.format(line))
            continue
        doc_info[index] = (data, number, name)
    return doc_info


def rename_docx(in_path: Path, out_path: Path, prefix: str, params) -> int:
    """Переименование документа согласно шаблону на основе выгрузки из базы"""
    number: str
    data: str
    data, number, _ = params
    file_name: str = in_path.name.split('.')[0].lstrip('t')
    file_name = str(100000 + int(file_name))[:6]
    number = number.replace('\\', ' ').replace('/', '')
    postfix: str = '{}={}={}={}.docx'.format(prefix, file_name, data, number)
    new_file: Path = Path(out_path).joinpath(postfix)
    try:
        in_path.replace(new_file)
        return 1
    except Exception as error:
        print(error)
        return 0


def validate_data_docx(data_file: tuple[str, str, str], path_doc: Path):
    """Проверка соответствия выгруженной информации по документу его содержимому"""
    doc: docx.Document = docx.Document(path_doc)
    fullText: list[str] = []
    for table in doc.tables:
        for cell in table._cells:
            fullText.append(cell.text)
    for paragraph in doc.paragraphs:
        fullText.append(paragraph.text)
    text_doc = ' '.join(fullText)
    status: str = ''
    data, number, name = data_file
    data = '.'.join(data.split('-')[::-1])
    number = number[:4].lstrip('0')
    for pattern in (data, number, name):
        if pattern in text_doc:
            status += '1'
        else:
            status += '0'
    return status


def save_json(path_file: Path, data):
    # Нужно ниже сделать аннотации, докстринг и рефокторинг
    import json
    path = path_file.name.split('.')[0]
    with open('{}.json'.format(path), 'w', encoding='utf-8') as outfile:
        json.dump(obj=data, fp=outfile,
                  ensure_ascii=False, indent=4,
                  sort_keys=True, separators=(',', ': '))


def validate_docx(path_file: Path, prefix: str, level: int = 1) -> None:
    """Валидация содержимого документов их данным согласно выгрузке из"""
    data = read_csv(path_file)
    save_json(path_file, data)
    # files = _iterdir(Path('to_validate'))
    # for file in files:
    #     index = int(file.name.split('.docx')[0].lstrip('t'))
    #     res = data.get(index, 'not')
    #     if res == 'not':
    #         print(index)
    #     elif type(res) == tuple:
    #         status = validate_data_docx(res, file)
    #         if (int(status[0]) + int(status[1]) + int(status[2])) >= level:
    #             if rename_docx(file, Path('output'), prefix, res):
    #                 del data[index]
    #         else:
    #             print('File {} {}'.format(index, status))
    #             print(res)
    # if data:
    #     answear = input('Остались документы в списке, показать? ')
    #     if answear.upper() in ('1', 'ДА', 'YES'):
    #         for key, val in data.items():
    #             print('{} - {}'.format(key, val))
    # else:
    #     print('Список документов пуст!')


if __name__ == '__main__':
    validate_docx(Path('D13.csv'), prefix='Л', level=1)
